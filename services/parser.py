import os
import uuid
import logging
import pdfplumber
import docx
from fastapi import UploadFile, HTTPException
from app.config import UPLOAD_DIR

logger = logging.getLogger("resume_analyzer.parser")

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber."""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error parsing PDF file {file_path}: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        doc = docx.Document(file_path)
        text_content = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
                
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Avoid duplicate cell texts if cell spans multiple columns
                    cell_text = cell.text.strip()
                    if cell_text and (not text_content or text_content[-1] != cell_text):
                        text_content.append(cell_text)
                        
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")

async def save_and_parse_file(file: UploadFile) -> tuple:
    """Saves the file to upload directory and returns (file_path, extracted_text)."""
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    
    if ext not in [".pdf", ".docx"]:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF or DOCX.")
    
    # Generate unique filename to avoid collision
    unique_filename = f"{uuid.uuid4()}{ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    # Save the file
    try:
        with open(file_path, "wb") as buffer:
            # Read in chunks to be memory efficient
            while True:
                chunk = await file.read(1024 * 1024) # 1MB chunks
                if not chunk:
                    break
                buffer.write(chunk)
    except Exception as e:
        logger.error(f"Failed to save file: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save uploaded file: {str(e)}")
        
    # Extract text
    extracted_text = ""
    if ext == ".pdf":
        extracted_text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        extracted_text = extract_text_from_docx(file_path)
        
    if not extracted_text.strip():
        # Clean up file if empty
        try:
            os.remove(file_path)
        except Exception:
            pass
        raise HTTPException(status_code=400, detail="The uploaded file contains no readable text.")
        
    return file_path, extracted_text
